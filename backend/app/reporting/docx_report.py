"""DOCX report builder (python-docx). Sections map 1:1 to TECHNICAL_GUIDE §8.1.

The single, deliberate departure from the sample: wherever an unverifiable
criterion cannot be proven, we print "⚠ Needs Manual Review" instead of a silent
"Pass". That honesty is the legitimacy upgrade.
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..models import Status
from ..storage import storage
from . import common

GLYPH = {
    Status.fail: ("● Fail", RGBColor(0xC0, 0x1A, 0x1A)),
    Status.partial: ("▲ Partial", RGBColor(0xB8, 0x86, 0x00)),
    Status.pass_: ("✔ Pass", RGBColor(0x1A, 0x7A, 0x2E)),
    Status.needs_manual_review: ("⚠ Needs Manual Review", RGBColor(0x8A, 0x6D, 0x00)),
    Status.not_applicable: ("— N/A", RGBColor(0x66, 0x66, 0x66)),
}
BRAND = RGBColor(0x0F, 0x3D, 0x57)


# ── low-level helpers ────────────────────────────────────────────────────────

def _status_run(paragraph, status: Status) -> None:
    label, color = GLYPH[status]
    run = paragraph.add_run(label)
    run.font.color.rgb = color
    run.bold = True


def _kv_table(doc, pairs: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in pairs:
        cells = t.add_row().cells
        kr = cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        cells[0].width = Inches(2.2)
        cells[1].text = str(v)
    doc.add_paragraph()


def _add_toc(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    for kind, txt in (("begin", None), ("instr", 'TOC \\o "1-3" \\h \\z \\u'),
                      ("separate", None), ("text", "Update this field (F9) to build the Table of Contents."),
                      ("end", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        elif kind == "text":
            el = OxmlElement("w:t")
            el.text = txt
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
    # Ask Word/LibreOffice to update fields on open.
    try:
        s = doc.settings.element
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        s.append(uf)
    except Exception:
        pass


def _embed_image(doc, key: str | None) -> None:
    if not key:
        p = doc.add_paragraph()
        r = p.add_run("Visual evidence was not captured for this finding "
                      "(element not locatable or screenshot engine unavailable).")
        r.italic = True
        return
    try:
        path = storage.local_path(key)
        doc.add_picture(str(path), width=Inches(6.0))
    except Exception as exc:  # noqa: BLE001
        p = doc.add_paragraph()
        p.add_run(f"[evidence image unavailable: {exc}]").italic = True


# ── sections ─────────────────────────────────────────────────────────────────

def _cover(doc, combo, meta) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Accessibility Compliance Audit Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"WCAG {combo['version']} — Level {combo['level']}")
    sr.font.size = Pt(16)

    for text in (meta["target_ref"], f"Scan date: {meta['scan_date'][:10]}",
                 meta["auditor_org"]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(text))
    doc.add_page_break()


def _document_control(doc, combo, meta) -> None:
    doc.add_heading("Document Control", level=1)
    _kv_table(doc, [
        ("Report Name", f"WCAG {combo['version']} Level {combo['level']} Audit"),
        ("Standard", f"WCAG {combo['version']} (Level {combo['level']}, cumulative)"),
        ("Target", meta["target_ref"]),
        ("Target Type", meta["target_type"]),
        ("Auditor", meta["auditor_org"]),
        ("Scan Date", meta["scan_date"][:19].replace("T", " ")),
        ("Version", "1.0"),
    ])


def _confidentiality(doc) -> None:
    doc.add_heading("Statement of Confidentiality and Limitations", level=1)
    doc.add_paragraph(
        "This report is confidential and intended solely for the named client. "
        "It reflects the state of the target at the time of scanning and the "
        "pages/screens actually reached.")
    doc.add_heading("Automated-testing limitations (read this)", level=2)
    for line in (
        "Automated tooling reliably detects only a minority of WCAG issues "
        "(~30–40%). A clean automated result is NOT a conformance claim.",
        "This platform prints \"Pass\" only for machine-decidable criteria that "
        "scanned clean with no engine error, or where a qualified human reviewer "
        "signed off. Every other criterion is marked \"Needs Manual Review\" — "
        "never a silent Pass.",
        "\"Needs Manual Review\" rows require a qualified human plus an "
        "assistive-technology pass (e.g. NVDA/VoiceOver) to close.",
        "Level AAA is largely manual by design; AAA reports are review-heavy.",
    ):
        doc.add_paragraph(line, style="List Bullet")


def _scope(doc, combo, meta) -> None:
    doc.add_heading("Scope", level=1)
    _kv_table(doc, [
        ("Activity", "Automated WCAG accessibility audit"),
        ("Target", meta["target_ref"]),
        ("Standard", f"WCAG {combo['version']} — Level {combo['level']}"),
        ("Hosts in scope (live)", ", ".join(meta.get("hosts_live", [])) or "—"),
        ("Pages crawled", str(meta.get("pages_crawled", 0))),
        ("Pages scanned (sampled)", str(meta.get("pages_scanned", 0))),
    ])


def _methodology(doc, meta) -> None:
    doc.add_heading("Testing Methodology", level=1)
    doc.add_paragraph(
        "Methodology follows the W3C WCAG-EM approach: define scope → explore → "
        "select a representative sample → audit → report. Discovery is passive "
        "and deny-by-default (only hosts on the authorized allowlist are touched).")
    doc.add_heading("Engines used", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(("Engine", "Clean", "Violations", "Errors")):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for eng, s in (meta.get("engines") or {}).items():
        c = t.add_row().cells
        c[0].text, c[1].text = eng, str(s.get("clean", 0))
        c[2].text, c[3].text = str(s.get("violations", 0)), str(s.get("error", 0))
    doc.add_paragraph()
    versions = meta.get("tool_versions") or {}
    doc.add_paragraph("Tool versions: " + ", ".join(f"{k} {v}" for k, v in versions.items()))


def _information_gathering(doc, meta) -> None:
    doc.add_heading("Information Gathering", level=1)
    _kv_table(doc, [
        ("Discovery sources", ", ".join(f"{k}={v}" for k, v in (meta.get("discovery") or {}).items()) or "—"),
        ("Templates identified", str(meta.get("templates", 0))),
        ("Out-of-scope hosts (not touched)", str(len(meta.get("hosts_out_of_scope", [])))),
    ])
    doc.add_paragraph(
        "Template sampling: pages sharing a structural template were clustered "
        "and a few representatives audited per template — accessibility defects "
        "recur per template, so this yields near-complete coverage efficiently.")


def _executive_summary(doc, combo) -> None:
    doc.add_heading("Executive Summary", level=1)
    cov = combo["coverage"]
    counts = combo["counts"]
    sev = combo["severity"]

    doc.add_heading("Coverage", level=2)
    _kv_table(doc, [
        ("Criteria in scope", str(cov["total_in_scope"])),
        ("Active (excluding obsolete)", str(cov["active_in_scope"])),
        ("Machine-decidable (auto)", str(cov["auto"])),
        ("Partly automatable (semi)", str(cov["semi"])),
        ("Manual-only", str(cov["manual"])),
    ])

    doc.add_heading("Results at a glance", level=2)
    _kv_table(doc, [(k, str(v)) for k, v in counts.items()])

    doc.add_heading("Confirmed-failure severity", level=2)
    _kv_table(doc, [(k.title(), str(v)) for k, v in sev.items()])

    rs = combo.get("review_stats", {})
    doc.add_heading("Human-review status", level=2)
    _kv_table(doc, [
        ("Manually verified (reviewer sign-off)", str(rs.get("manually_verified", 0))),
        ("Open — awaiting manual review", str(rs.get("open_review", 0))),
    ])


def _checklist(doc, combo) -> None:
    doc.add_heading("Full WCAG Compliance Checklist", level=1)
    rows = combo["rows"]
    by_principle: dict[str, list[dict]] = {}
    for r in rows:
        by_principle.setdefault(r["crit"]["principle"], []).append(r)

    for principle, prows in by_principle.items():
        doc.add_heading(principle, level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for i, h in enumerate(("SC", "Name", "Level", "Status", "Note")):
            t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for r in sorted(prows, key=lambda x: [int(n) for n in x["crit"]["num"].split(".")]):
            c = t.add_row().cells
            c[0].text = r["crit"]["num"]
            c[1].text = r["crit"]["name"]
            c[2].text = r["crit"]["level"]
            _status_run(c[3].paragraphs[0], r["status"])
            c[4].text = r["note"]


def _detailed_observations(doc, combo) -> None:
    doc.add_heading("Detailed Observations", level=1)
    if not combo["issues"]:
        doc.add_paragraph("No automated failures were confirmed for this "
                          "combination. Remaining criteria require manual review "
                          "(see the checklist).")
        return
    for i, f in enumerate(combo["issues"], 1):
        doc.add_heading(f"Issue {i}: {f.sc_name} ({f.sc_num})", level=2)
        locs = "\n".join(f"{loc.ref} ({loc.count})" for loc in f.locations)
        _kv_table(doc, [
            ("WCAG Criterion", f"{f.sc_num} {f.sc_name}"),
            ("Conformance Level", f.level),
            ("Impact Rating", f.impact.value.title()),
            ("Detected by", ", ".join(f.engines_agreeing) or f.engine),
            ("Confidence", f.confidence.value),
            ("Total Occurrences", str(f.occurrences)),
            ("Element Location(s)", locs),
            ("Rule", f"{f.engine_rule_id}"),
            ("Description", f.description or "—"),
        ])
        if f.computed:
            _kv_table(doc, [("Measured", ", ".join(f"{k}={v}" for k, v in list(f.computed.items())[:8]))])
        doc.add_paragraph().add_run("Visual Evidence").bold = True
        _embed_image(doc, f.screenshot_key)
        doc.add_paragraph().add_run("Remediation").bold = True
        doc.add_paragraph(f.remediation)
        doc.add_paragraph()


def _risk_and_recommendations(doc, combo) -> None:
    doc.add_heading("Risk Evaluation and Conclusion", level=1)
    sev = combo["severity"]
    total = sum(sev.values())
    doc.add_paragraph(
        f"{total} confirmed automated failure group(s): "
        f"{sev['blocker']} blocker, {sev['serious']} serious, "
        f"{sev['moderate']} moderate, {sev['minor']} minor. "
        f"{combo['counts']['Needs Manual Review']} criteria still require manual "
        "review before a conformance determination can be made.")

    doc.add_heading("Recommendations (prioritized)", level=1)
    order = {"blocker": "P1", "serious": "P1", "moderate": "P2", "minor": "P3"}
    seen = set()
    for f in sorted(combo["issues"], key=lambda x: x.impact.value):
        if f.sc_num in seen:
            continue
        seen.add(f.sc_num)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{order[f.impact.value]}] {f.sc_num} {f.sc_name}: ").bold = True
        p.add_run(f.remediation)
    doc.add_paragraph(
        "Then complete the manual-review pass (assistive technology) to close the "
        "\"Needs Manual Review\" criteria.")


def _prepared_by(doc, meta) -> None:
    doc.add_heading("Prepared By", level=1)
    doc.add_paragraph(meta["auditor_org"])
    doc.add_paragraph(f"Generated {meta['scan_date'][:19].replace('T', ' ')} UTC")

    doc.add_heading("Reproducibility", level=2)
    versions = meta.get("tool_versions") or {}
    doc.add_paragraph(
        "Results reflect the target's state at scan time and the pages/screens "
        "actually reached. Engine & tool versions used:")
    for name, ver in versions.items():
        doc.add_paragraph(f"{name}: {ver}", style="List Bullet")
    doc.add_paragraph(
        "Testability classification (auto/semi/manual) is this platform's editorial "
        "engineering judgment on machine-decidability — it is not part of the W3C "
        "standard. Raw engine output is retained for every finding for traceability.")


# ── entry point ──────────────────────────────────────────────────────────────

def build_docx(combo: dict, meta: dict, out_path: str) -> str:
    doc = Document()
    _cover(doc, combo, meta)
    doc.add_heading("Table of Contents", level=1)
    _add_toc(doc)
    doc.add_page_break()
    _document_control(doc, combo, meta)
    _confidentiality(doc)
    _scope(doc, combo, meta)
    _methodology(doc, meta)
    _information_gathering(doc, meta)
    _executive_summary(doc, combo)
    _checklist(doc, combo)
    _detailed_observations(doc, combo)
    _risk_and_recommendations(doc, combo)
    _prepared_by(doc, meta)
    doc.save(out_path)
    return out_path
