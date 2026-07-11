"""VPAT 2.5 Accessibility Conformance Report (ACR) export.

Maps the platform's internal statuses onto the five ITI VPAT/ACR conformance
terms and renders a WCAG-edition ACR (.docx). VPAT 2.5 (2025-04-24) is the first
edition to cover WCAG 2.2.

The mapping preserves the anti-fabrication rule: an unproven criterion becomes
**Not Evaluated** (never silently "Supports").
"""
from __future__ import annotations

from docx import Document

from ..models import Status

# Platform status → ACR term.
ACR_TERM = {
    Status.pass_: "Supports",
    Status.partial: "Partially Supports",
    Status.fail: "Does Not Support",
    Status.not_applicable: "Not Applicable",
    Status.needs_manual_review: "Not Evaluated",
}

TERMS_GLOSS = [
    ("Supports", "The functionality meets the criterion without exceptions."),
    ("Partially Supports", "Some functionality does not meet the criterion."),
    ("Does Not Support", "Most functionality does not meet the criterion."),
    ("Not Applicable", "The criterion is not relevant to the product."),
    ("Not Evaluated", "Not verified — this platform only claims conformance for "
                      "machine-decidable criteria that scanned clean, or where a "
                      "reviewer signed off; all else is Not Evaluated."),
]


def acr_term(status: Status) -> str:
    return ACR_TERM.get(status, "Not Evaluated")


def _remark(row: dict) -> str:
    st = row["status"]
    if st == Status.fail:
        return row["crit"].get("_issue_summary", "") or "Automated failure detected — see the detailed report."
    if row.get("via_review"):
        return row["note"]
    if st == Status.needs_manual_review:
        return "Requires manual / assistive-technology review."
    if st == Status.not_applicable:
        return row["note"] or "Not applicable."
    return row["note"]


def generate_acr(combo: dict, meta: dict, out_path: str) -> str:
    doc = Document()
    doc.add_heading("Accessibility Conformance Report", level=0)
    doc.add_paragraph(f"VPAT® 2.5 (WCAG edition) — WCAG {combo['version']} "
                      f"Level {combo['level']}")

    doc.add_heading("Product information", level=1)
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in [("Name of Product", meta["target_ref"]),
                 ("Report Date", meta["scan_date"][:10]),
                 ("Evaluation Method", "Automated multi-engine scan"
                  + (" + manual review" if combo.get("review_stats", {}).get("manually_verified") else "")),
                 ("Evaluator", meta["auditor_org"]),
                 ("Standard", f"WCAG {combo['version']} (Level {combo['level']})")]:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(k).bold = True
        c[1].text = str(v)

    doc.add_heading("Conformance terms", level=1)
    for term, gloss in TERMS_GLOSS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{term}: ").bold = True
        p.add_run(gloss)

    # One table per WCAG level present in scope (VPAT Table 1/2/3 convention).
    levels = {"A": [], "AA": [], "AAA": []}
    for row in combo["rows"]:
        levels[row["crit"]["level"]].append(row)

    for lvl in ("A", "AA", "AAA"):
        rows = levels[lvl]
        if not rows:
            continue
        doc.add_heading(f"Success Criteria, Level {lvl}", level=1)
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
        for i, h in enumerate(("Criteria", "Conformance Level", "Remarks and Explanations")):
            tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for row in sorted(rows, key=lambda r: [int(n) for n in r["crit"]["num"].split(".")]):
            c = tbl.add_row().cells
            c[0].text = f"{row['crit']['num']} {row['crit']['name']}"
            c[1].text = acr_term(row["status"])
            c[2].text = _remark(row)

    doc.save(out_path)
    return out_path
